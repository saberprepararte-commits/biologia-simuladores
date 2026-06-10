from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Guia_estudiante_osmosis_indagacion.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "AFC4D8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def width_to_inches(dxa):
    return Inches(dxa / 1440)


def style_paragraph(paragraph, size=11, color=INK, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = bold
        run.font.italic = italic


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(BLUE)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(INK)
        run.font.size = Pt(11)
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    style_paragraph(paragraph)
    return paragraph


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(11)
    p.add_run(" " + text)
    style_paragraph(p)
    return table


def add_blank_lines(cell, count=2):
    for _ in range(count):
        p = cell.add_paragraph("_" * 62)
        p.paragraph_format.space_after = Pt(2)
        style_paragraph(p, size=10, color="555555")


def fill_header_row(row, labels):
    for cell, label in zip(row.cells, labels):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label)
        style_paragraph(p, size=10, color=INK, bold=True)


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Guía de indagación: ósmosis"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(header, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.text = "Nombre: ____________________   Curso: ______   Fecha: __________"
    style_paragraph(footer, size=9, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Guía de trabajo del estudiante")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Ósmosis y transporte de agua a través de una membrana semipermeable")
    style_paragraph(subtitle, size=12, color=DARK_BLUE, bold=True)

    meta = doc.add_table(rows=2, cols=4)
    set_table_widths(meta, [1800, 2880, 1800, 2880])
    values = [
        ["Estudiante", "", "Grupo", ""],
        ["Fecha", "", "Tiempo estimado", "45-60 minutos"],
    ]
    for row, row_values in zip(meta.rows, values):
        for i, cell in enumerate(row.cells):
            cell.text = row_values[i]
            if i % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                style_paragraph(cell.paragraphs[0], size=10, bold=True)
            else:
                style_paragraph(cell.paragraphs[0], size=10)

    add_callout(
        doc,
        "Propósito de la guía:",
        "desarrollar la competencia de indagación científica usando el simulador de ósmosis. "
        "Vas a formular una pregunta, proponer una hipótesis, controlar variables, recoger datos, "
        "analizar evidencias y construir una conclusión basada en resultados."
    )

    add_heading(doc, "1. Punto de partida: observar antes de explicar", 1)
    add_body(
        doc,
        "Explora el simulador sin presionar Start durante un minuto. Observa la célula, las moléculas de agua, "
        "los solutos y los controles de concentración."
    )
    obs = doc.add_table(rows=4, cols=2)
    set_table_widths(obs, [2500, 6860])
    fill_header_row(obs.rows[0], ["Aspecto observado", "Registro del estudiante"])
    obs.cell(1, 0).text = "¿Qué cambia cuando modificas las barras de soluto?"
    obs.cell(2, 0).text = "¿Qué representa el agua dentro y fuera de la célula?"
    obs.cell(3, 0).text = "¿Qué representa el soluto y por qué no atraviesa la membrana?"
    for row in obs.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "2. Pregunta investigable", 1)
    add_body(
        doc,
        "Una pregunta investigable relaciona variables y puede responderse con datos del simulador. "
        "Evita preguntas que solo pidan una definición."
    )
    qtab = doc.add_table(rows=3, cols=2)
    set_table_widths(qtab, [2600, 6760])
    fill_header_row(qtab.rows[0], ["Criterio", "Tu propuesta"])
    qtab.cell(1, 0).text = "Pregunta investigable"
    qtab.cell(2, 0).text = "¿Qué relación esperas estudiar?"
    for row in qtab.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "3. Hipótesis y variables", 1)
    add_body(
        doc,
        "Escribe una hipótesis con la estructura: Si..., entonces..., porque... Debe anticipar el movimiento "
        "del agua según la concentración de soluto."
    )
    hyp = doc.add_table(rows=5, cols=2)
    set_table_widths(hyp, [2600, 6760])
    fill_header_row(hyp.rows[0], ["Elemento de indagación", "Respuesta"])
    labels = [
        "Hipótesis",
        "Variable independiente",
        "Variable dependiente",
        "Variables controladas",
    ]
    for idx, label in enumerate(labels, start=1):
        hyp.cell(idx, 0).text = label
        style_paragraph(hyp.cell(idx, 0).paragraphs[0], size=10, bold=True)
        add_blank_lines(hyp.cell(idx, 1), 2 if idx in (1, 4) else 1)

    add_heading(doc, "4. Diseño experimental en el simulador", 1)
    add_body(
        doc,
        "Realiza al menos cinco ensayos. En cada ensayo cambia las concentraciones, predice qué pasará, presiona "
        "Start, observa hasta que se acerque al equilibrio y registra evidencia."
    )
    design = doc.add_table(rows=7, cols=7)
    set_table_widths(design, [700, 1180, 1180, 1480, 1320, 1480, 2020])
    fill_header_row(
        design.rows[0],
        ["Ensayo", "Soluto dentro (%)", "Soluto fuera (%)", "Predicción", "Dirección del agua", "Cambio celular", "Evidencia"],
    )
    presets = [
        ["1", "50", "50", "", "", "", ""],
        ["2", "25", "75", "", "", "", ""],
        ["3", "75", "25", "", "", "", ""],
        ["4", "", "", "", "", "", ""],
        ["5", "", "", "", "", "", ""],
        ["6 opcional", "", "", "", "", "", ""],
    ]
    for row, values in zip(design.rows[1:], presets):
        for cell, value in zip(row.cells, values):
            cell.text = value
            style_paragraph(cell.paragraphs[0], size=9)
    add_body(
        doc,
        "Consejo de control experimental: cambia solo una condición a la vez cuando quieras comparar resultados. "
        "Usa Restart antes de iniciar un nuevo ensayo."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "5. Análisis de datos", 1)
    add_body(
        doc,
        "Ahora transforma tus observaciones en explicación científica. Usa los datos de la tabla anterior, no solo "
        "lo que esperabas que ocurriera."
    )
    analysis = doc.add_table(rows=5, cols=2)
    set_table_widths(analysis, [3000, 6360])
    fill_header_row(analysis.rows[0], ["Pregunta de análisis", "Respuesta con evidencia"])
    prompts = [
        "¿En qué ensayos el agua entró principalmente a la célula?",
        "¿En qué ensayos el agua salió principalmente de la célula?",
        "¿Qué condición produjo equilibrio o movimiento neto cercano a cero?",
        "¿Qué patrón general encontraste entre concentración de soluto y movimiento del agua?",
    ]
    for row, prompt in zip(analysis.rows[1:], prompts):
        row.cells[0].text = prompt
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "6. Construcción de explicación: modelo CER", 1)
    add_body(
        doc,
        "CER significa afirmación, evidencia y razonamiento. Úsalo para explicar tus resultados con lenguaje científico."
    )
    cer = doc.add_table(rows=4, cols=2)
    set_table_widths(cer, [2200, 7160])
    fill_header_row(cer.rows[0], ["Parte", "Escritura del estudiante"])
    cer_parts = [
        ("Afirmación", "Responde directamente tu pregunta investigable."),
        ("Evidencia", "Incluye datos concretos de dos o más ensayos."),
        ("Razonamiento", "Explica por qué el agua se mueve hacia donde hay mayor concentración de soluto."),
    ]
    for row, (label, hint) in zip(cer.rows[1:], cer_parts):
        row.cells[0].text = f"{label}\n{hint}"
        style_paragraph(row.cells[0].paragraphs[0], size=9, bold=True)
        add_blank_lines(row.cells[1], 3)

    add_heading(doc, "7. Evaluación de la indagación", 1)
    rubric = doc.add_table(rows=5, cols=4)
    set_table_widths(rubric, [2100, 2420, 2420, 2420])
    fill_header_row(rubric.rows[0], ["Criterio", "Logrado", "En proceso", "Por mejorar"])
    criteria = [
        "Formulé una pregunta investigable que relaciona variables.",
        "Identifiqué variables y controles del experimento.",
        "Registré datos suficientes y observaciones claras.",
        "Construí una conclusión con evidencia y razonamiento.",
    ]
    for row, criterion in zip(rubric.rows[1:], criteria):
        row.cells[0].text = criterion
        style_paragraph(row.cells[0].paragraphs[0], size=9, bold=True)
        for cell in row.cells[1:]:
            cell.text = "☐"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(cell.paragraphs[0], size=14)

    add_heading(doc, "8. Reflexión final", 1)
    refl = doc.add_table(rows=3, cols=2)
    set_table_widths(refl, [3000, 6360])
    fill_header_row(refl.rows[0], ["Reflexiona", "Respuesta"])
    refl.cell(1, 0).text = "¿Qué cambiarías si repitieras la indagación?"
    refl.cell(2, 0).text = "¿Cómo se relaciona la ósmosis con células reales del cuerpo?"
    for row in refl.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 3)

    doc.save(OUT)


if __name__ == "__main__":
    make_doc()
