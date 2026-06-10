from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_student_guide import (
    BLUE,
    DARK_BLUE,
    INK,
    LIGHT_BLUE,
    add_blank_lines,
    add_body,
    add_callout,
    add_heading,
    fill_header_row,
    set_table_widths,
    style_paragraph,
)


OUT = "Student_Worksheet_Osmosis_Inquiry.docx"


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Inquiry worksheet: osmosis"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(header, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.text = "Name: ____________________   Class: ______   Date: __________"
    style_paragraph(footer, size=9, color="666666")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Student Worksheet")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Osmosis and water transport across a semipermeable membrane")
    style_paragraph(subtitle, size=12, color=DARK_BLUE, bold=True)

    meta = doc.add_table(rows=2, cols=4)
    set_table_widths(meta, [1800, 2880, 1800, 2880])
    values = [
        ["Student", "", "Group", ""],
        ["Date", "", "Estimated time", "45-60 minutes"],
    ]
    for row, row_values in zip(meta.rows, values):
        for i, cell in enumerate(row.cells):
            cell.text = row_values[i]
            if i % 2 == 0:
                from build_student_guide import set_cell_shading
                set_cell_shading(cell, LIGHT_BLUE)
                style_paragraph(cell.paragraphs[0], size=10, bold=True)
            else:
                style_paragraph(cell.paragraphs[0], size=10)

    add_callout(
        doc,
        "Purpose of this worksheet:",
        "to develop scientific inquiry skills using the osmosis simulator. You will formulate a question, "
        "propose a hypothesis, control variables, collect data, analyze evidence, and write a conclusion "
        "based on results."
    )

    add_heading(doc, "1. Starting point: observe before explaining", 1)
    add_body(
        doc,
        "Explore the simulator for one minute without pressing Start. Observe the cell, water molecules, "
        "solute particles, and concentration controls."
    )
    obs = doc.add_table(rows=4, cols=2)
    set_table_widths(obs, [2500, 6860])
    fill_header_row(obs.rows[0], ["Observed aspect", "Student record"])
    obs.cell(1, 0).text = "What changes when you move the solute sliders?"
    obs.cell(2, 0).text = "What does the water inside and outside the cell represent?"
    obs.cell(3, 0).text = "What does the solute represent, and why does it not cross the membrane?"
    for row in obs.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "2. Investigable question", 1)
    add_body(
        doc,
        "An investigable question connects variables and can be answered with data from the simulator. "
        "Avoid questions that only ask for a definition."
    )
    qtab = doc.add_table(rows=3, cols=2)
    set_table_widths(qtab, [2600, 6760])
    fill_header_row(qtab.rows[0], ["Criterion", "Your proposal"])
    qtab.cell(1, 0).text = "Investigable question"
    qtab.cell(2, 0).text = "What relationship do you expect to study?"
    for row in qtab.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "3. Hypothesis and variables", 1)
    add_body(
        doc,
        "Write a hypothesis using this structure: If..., then..., because... It should predict water movement "
        "according to solute concentration."
    )
    hyp = doc.add_table(rows=5, cols=2)
    set_table_widths(hyp, [2600, 6760])
    fill_header_row(hyp.rows[0], ["Inquiry element", "Response"])
    labels = [
        "Hypothesis",
        "Independent variable",
        "Dependent variable",
        "Controlled variables",
    ]
    for idx, label in enumerate(labels, start=1):
        hyp.cell(idx, 0).text = label
        style_paragraph(hyp.cell(idx, 0).paragraphs[0], size=10, bold=True)
        add_blank_lines(hyp.cell(idx, 1), 2 if idx in (1, 4) else 1)

    add_heading(doc, "4. Experimental design in the simulator", 1)
    add_body(
        doc,
        "Run at least five trials. In each trial, change the concentrations, predict what will happen, press "
        "Start, observe until the system approaches equilibrium, and record evidence."
    )
    design = doc.add_table(rows=7, cols=7)
    set_table_widths(design, [700, 1180, 1180, 1480, 1320, 1480, 2020])
    fill_header_row(
        design.rows[0],
        ["Trial", "Solute inside (%)", "Solute outside (%)", "Prediction", "Water direction", "Cell change", "Evidence"],
    )
    presets = [
        ["1", "50", "50", "", "", "", ""],
        ["2", "25", "75", "", "", "", ""],
        ["3", "75", "25", "", "", "", ""],
        ["4", "", "", "", "", "", ""],
        ["5", "", "", "", "", "", ""],
        ["6 optional", "", "", "", "", "", ""],
    ]
    for row, values in zip(design.rows[1:], presets):
        for cell, value in zip(row.cells, values):
            cell.text = value
            style_paragraph(cell.paragraphs[0], size=9)
    add_body(
        doc,
        "Experimental control tip: change only one condition at a time when you want to compare results. "
        "Use Restart before beginning a new trial."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "5. Data analysis", 1)
    add_body(
        doc,
        "Now turn your observations into a scientific explanation. Use the data from the previous table, "
        "not only what you expected to happen."
    )
    analysis = doc.add_table(rows=5, cols=2)
    set_table_widths(analysis, [3000, 6360])
    fill_header_row(analysis.rows[0], ["Analysis question", "Evidence-based response"])
    prompts = [
        "In which trials did water mainly enter the cell?",
        "In which trials did water mainly leave the cell?",
        "Which condition produced equilibrium or nearly zero net movement?",
        "What general pattern did you find between solute concentration and water movement?",
    ]
    for row, prompt in zip(analysis.rows[1:], prompts):
        row.cells[0].text = prompt
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 2)

    add_heading(doc, "6. Building an explanation: CER model", 1)
    add_body(
        doc,
        "CER means claim, evidence, and reasoning. Use it to explain your results with scientific language."
    )
    cer = doc.add_table(rows=4, cols=2)
    set_table_widths(cer, [2200, 7160])
    fill_header_row(cer.rows[0], ["Part", "Student writing"])
    cer_parts = [
        ("Claim", "Directly answer your investigable question."),
        ("Evidence", "Include specific data from two or more trials."),
        ("Reasoning", "Explain why water moves toward the area with higher solute concentration."),
    ]
    for row, (label, hint) in zip(cer.rows[1:], cer_parts):
        row.cells[0].text = f"{label}\n{hint}"
        style_paragraph(row.cells[0].paragraphs[0], size=9, bold=True)
        add_blank_lines(row.cells[1], 3)

    add_heading(doc, "7. Inquiry self-assessment", 1)
    rubric = doc.add_table(rows=5, cols=4)
    set_table_widths(rubric, [2100, 2420, 2420, 2420])
    fill_header_row(rubric.rows[0], ["Criterion", "Achieved", "In progress", "Needs improvement"])
    criteria = [
        "I formulated an investigable question that connects variables.",
        "I identified variables and controls in the experiment.",
        "I recorded enough data and clear observations.",
        "I wrote a conclusion with evidence and reasoning.",
    ]
    for row, criterion in zip(rubric.rows[1:], criteria):
        row.cells[0].text = criterion
        style_paragraph(row.cells[0].paragraphs[0], size=9, bold=True)
        for cell in row.cells[1:]:
            cell.text = "☐"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(cell.paragraphs[0], size=14)

    add_heading(doc, "8. Final reflection", 1)
    refl = doc.add_table(rows=3, cols=2)
    set_table_widths(refl, [3000, 6360])
    fill_header_row(refl.rows[0], ["Reflect", "Response"])
    refl.cell(1, 0).text = "What would you change if you repeated the inquiry?"
    refl.cell(2, 0).text = "How is osmosis related to real cells in the body?"
    for row in refl.rows[1:]:
        style_paragraph(row.cells[0].paragraphs[0], size=10, bold=True)
        add_blank_lines(row.cells[1], 3)

    doc.save(OUT)


if __name__ == "__main__":
    make_doc()
