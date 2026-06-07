from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "guia-assets"
OUT = ROOT / "Guia_estudiante_ciclo_nitrogeno.docx"

COLORS = {
    "navy": "083248",
    "deep": "0A4A5A",
    "teal": "0E99A8",
    "cyan": "28B5DC",
    "mint": "BDEEDB",
    "green": "2D9658",
    "gold": "D98B32",
    "red": "C84B59",
    "cream": "FFF8E8",
    "ice": "EAF5F6",
    "soil": "704526",
    "white": "FFFFFF",
}


def crop_image(src, dest, box):
    image = Image.open(src)
    image.crop(box).save(dest)


def prepare_images():
    control = ASSETS / "captura-control.png"
    toxic = ASSETS / "captura-toxicidad.png"
    crop_image(control, ASSETS / "recorte-escena-control.png", (285, 185, 1285, 760))
    crop_image(control, ASSETS / "recorte-panel-control.png", (0, 80, 280, 995))
    crop_image(control, ASSETS / "recorte-lecturas-control.png", (1300, 375, 1585, 775))
    crop_image(toxic, ASSETS / "recorte-toxicidad-lecturas.png", (1300, 375, 1585, 710))
    crop_image(toxic, ASSETS / "recorte-toxicidad-consecuencia.png", (1300, 720, 1585, 980))
    crop_image(toxic, ASSETS / "recorte-toxicidad-centro.png", (285, 185, 1285, 760))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="FFFFFF", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="083248", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_colored_table(doc, headers, rows, widths=None, header_fill="0E99A8", band_fill="EAF5F6"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=8)
        if widths:
            cell.width = widths[i]
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        fill = band_fill if row_index % 2 == 0 else "FFFFFF"
        for i, value in enumerate(row):
            set_cell_shading(cells[i], fill)
            set_cell_border(cells[i], "D4E5E8", "6")
            set_cell_text(cells[i], str(value), color="083248", size=8)
            if widths:
                cells[i].width = widths[i]
    return table


def add_title(doc, text, level=1, color="083248"):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = True
    return paragraph


def add_callout(doc, title, body, fill="EAF5F6", border="0E99A8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, "16")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Aptos"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(COLORS["deep"])
    return table


def add_body(doc, text, size=10, bold=False, color="12324A", align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Ciclo del nitrógeno BioLab | Guía de trabajo experimental")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("5D7480")


def build_doc():
    prepare_images()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Aptos Display"
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor.from_string(COLORS["navy"])

    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover.cell(0, 0)
    set_cell_shading(cell, COLORS["navy"])
    set_cell_border(cell, COLORS["navy"], "6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guía de trabajo del estudiante")
    r.font.name = "Aptos Display"
    r.font.size = Pt(23)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(COLORS["white"])
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Laboratorio experimental: ciclo del nitrógeno")
    r2.font.name = "Aptos"
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor.from_string("9DE9F7")

    doc.add_picture(str(ASSETS / "portada-guia-biolab.png"), width=Inches(7.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(
        doc,
        "Tu misión",
        "Vas a modificar variables del simulador, registrar evidencias y explicar las consecuencias ecológicas. No se trata de adivinar: se trata de probar, comparar y defender conclusiones con datos.",
        fill="E2F7FA",
    )
    meta = add_colored_table(
        doc,
        ["Nombre", "Curso", "Fecha", "Equipo"],
        [["", "", "", ""]],
        header_fill=COLORS["green"],
        band_fill="FFFFFF",
    )
    for row in meta.rows:
        for cell in row.cells:
            cell.height = Cm(0.8)
    doc.add_page_break()

    add_title(doc, "1. Punto de partida", 1)
    add_body(
        doc,
        "El ciclo del nitrógeno depende de entradas, transformaciones y pérdidas. En el simulador observarás cómo cambian los indicadores cuando alteras plantas, bacterias fijadoras, amoníaco/amonio y nitrato.",
    )
    add_callout(
        doc,
        "Pregunta problema",
        "¿Qué combinación de variables mantiene un ciclo equilibrado y en qué condiciones aparecen toxicidad, baja fijación o pérdida de nitrato por lixiviación?",
        fill="FFF2D9",
        border=COLORS["gold"],
    )
    add_title(doc, "Variables del laboratorio", 2)
    add_colored_table(
        doc,
        ["Variable", "Qué representa", "Consecuencia esperada al cambiarla"],
        [
            ["Número de plantas", "Demanda de nitrato y asimilación vegetal.", "Muchas plantas aumentan la demanda; pocas plantas pueden dejar nitrato sin usar."],
            ["Bacterias fijadoras de N₂", "Entrada biológica de nitrógeno al suelo.", "Pocas reducen la fijación; muchas elevan la entrada de compuestos nitrogenados."],
            ["Amoníaco / amonio", "Sustrato para nitrificación.", "En rango moderado ayuda; en exceso puede causar toxicidad en raíces y microorganismos."],
            ["Nitrato", "Forma asimilable por las plantas.", "Si falta limita crecimiento; si sobra aumenta lixiviación y desnitrificación."],
        ],
        header_fill=COLORS["teal"],
    )
    doc.add_paragraph()
    add_title(doc, "Grupo control", 2)
    doc.add_picture(str(ASSETS / "recorte-escena-control.png"), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Observa el grupo control antes de tocar los deslizadores. Ese será tu punto de comparación para todos los tratamientos.", size=9)
    doc.add_page_break()

    add_title(doc, "2. Método experimental", 1)
    add_body(
        doc,
        "Trabaja como en un laboratorio: cambia una variable a la vez, registra los indicadores y compara contra el grupo control. Cuando pruebes combinaciones, explica por qué cambiaste más de una variable.",
    )
    steps = [
        ["1", "Restablece el grupo control.", "Anota los valores iniciales de plantas, bacterias, amoníaco y nitrato."],
        ["2", "Formula una hipótesis.", "Predice qué indicador subirá o bajará y justifica tu predicción."],
        ["3", "Modifica una variable.", "Mueve el deslizador lentamente y observa barras, moléculas y texto de consecuencias."],
        ["4", "Registra datos.", "Copia los valores finales y describe la evidencia visual del simulador."],
        ["5", "Interpreta.", "Relaciona los resultados con fijación, nitrificación, asimilación, lixiviación o toxicidad."],
    ]
    add_colored_table(doc, ["Paso", "Acción", "Evidencia que debes recoger"], steps, header_fill=COLORS["green"], band_fill="EAF8F0")
    add_callout(
        doc,
        "Regla de oro",
        "Si cambias muchas variables al mismo tiempo, es más difícil saber cuál causó el efecto. Primero aísla; después combina.",
        fill="EAF5F6",
    )
    add_title(doc, "Panel y lecturas", 2)
    panel_table = doc.add_table(rows=1, cols=2)
    panel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = panel_table.rows[0].cells
    set_cell_shading(left, "FFFFFF")
    set_cell_shading(right, "FFFFFF")
    set_cell_border(left, "D4E5E8")
    set_cell_border(right, "D4E5E8")
    left.paragraphs[0].add_run().add_picture(str(ASSETS / "recorte-panel-control.png"), width=Inches(2.0))
    right.paragraphs[0].add_run().add_picture(str(ASSETS / "recorte-lecturas-control.png"), width=Inches(2.8))
    doc.add_page_break()

    add_title(doc, "3. Registro de experimentos", 1)
    add_body(doc, "Completa una fila por cada tratamiento. Usa números aproximados si el simulador no muestra decimales.")
    headers = [
        "Tratamiento",
        "Plantas",
        "Bacterias",
        "NH₄⁺",
        "NO₃⁻",
        "Disp.",
        "Fij.",
        "Lix.",
        "Estrés",
        "Balance",
    ]
    rows = [["Control", "", "", "", "", "", "", "", "", ""]]
    rows += [[f"Prueba {i}", "", "", "", "", "", "", "", "", ""] for i in range(1, 7)]
    add_colored_table(doc, headers, rows, header_fill=COLORS["navy"], band_fill="F2FBFC")
    add_body(doc, "Disp. = disponibilidad para plantas | Fij. = fijación biológica | Lix. = riesgo de lixiviación.", size=8, color="4D6670")
    add_title(doc, "Observaciones cualitativas", 2)
    add_colored_table(
        doc,
        ["Prueba", "¿Qué cambió en las moléculas o flechas?", "¿Qué consecuencia mostró el simulador?", "Explicación biológica"],
        [[f"{i}", "", "", ""] for i in range(1, 5)],
        header_fill=COLORS["gold"],
        band_fill="FFF8E8",
    )
    doc.add_page_break()

    add_title(doc, "4. Experimentos guiados", 1)
    experiments = [
        ["A. Amoníaco bajo", "Baja el amoníaco cerca de 0 mg/kg.", "La nitrificación tendrá menos sustrato; puede bajar el nitrato disponible."],
        ["B. Amoníaco tóxico", "Sube el amoníaco al máximo.", "Aparece estrés por amoníaco; raíces y microbios pueden verse afectados."],
        ["C. Nitrato escaso", "Baja el nitrato y conserva plantas altas.", "La asimilación vegetal queda limitada aunque haya demanda."],
        ["D. Nitrato excesivo", "Sube el nitrato por encima de la demanda vegetal.", "Aumenta el riesgo de lixiviación y pérdida de fertilidad."],
        ["E. Pocas bacterias fijadoras", "Reduce bacterias fijadoras de N₂.", "Disminuye la entrada biológica de nitrógeno al suelo."],
        ["F. Sistema equilibrado", "Busca una combinación con balance alto y bajo estrés.", "Justifica por qué tu combinación se aproxima a un suelo estable."],
    ]
    add_colored_table(doc, ["Experimento", "Manipulación", "Predicción que debes comprobar"], experiments, header_fill=COLORS["teal"], band_fill="E8F8FA")
    add_title(doc, "Escenario de desequilibrio", 2)
    doc.add_picture(str(ASSETS / "recorte-toxicidad-centro.png"), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_callout(
        doc,
        "Lectura clave",
        "Cuando amoníaco y nitrato están muy altos, no basta con decir que hay más nutrientes. También debes evaluar toxicidad, actividad microbiana, lixiviación y balance del ciclo.",
        fill="FBE7EA",
        border=COLORS["red"],
    )
    doc.add_page_break()

    add_title(doc, "5. Análisis de resultados", 1)
    prompts = [
        ["Relación causa-efecto", "Elige dos pruebas y explica qué variable causó el mayor cambio en el balance del ciclo."],
        ["Amoníaco", "¿Por qué una concentración moderada puede ayudar, pero una concentración alta puede ser perjudicial?"],
        ["Nitrato", "¿En qué momento el nitrato deja de ser ventaja para convertirse en riesgo de pérdida?"],
        ["Lixiviación baja", "¿Cuándo conservar nitrato en el suelo es positivo y cuándo puede indicar acumulación?"],
        ["Lixiviación alta", "¿Qué consecuencias tendría para fertilidad del suelo y aguas subterráneas?"],
        ["Microorganismos", "¿Qué pasaría con el ciclo si disminuye la actividad microbiana? Usa evidencia del simulador."],
    ]
    add_colored_table(doc, ["Idea a analizar", "Respuesta con evidencia"], prompts, header_fill=COLORS["green"], band_fill="F1FAF4")
    add_title(doc, "Comparación control vs. desequilibrio", 2)
    two = doc.add_table(rows=1, cols=2)
    two.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = two.rows[0].cells
    set_cell_shading(c1, "EAF8F0")
    set_cell_shading(c2, "FBE7EA")
    set_cell_border(c1, COLORS["green"])
    set_cell_border(c2, COLORS["red"])
    c1.paragraphs[0].add_run("Control estable").bold = True
    c1.add_paragraph().add_run().add_picture(str(ASSETS / "recorte-lecturas-control.png"), width=Inches(2.1))
    c2.paragraphs[0].add_run("Amoníaco y nitrato altos").bold = True
    c2.add_paragraph().add_run().add_picture(str(ASSETS / "recorte-toxicidad-lecturas.png"), width=Inches(2.1))
    add_callout(
        doc,
        "Consecuencia esperada",
        "El amoníaco muy alto puede reducir el crecimiento de raíces y la actividad microbiana. El nitrato excesivo aumenta el riesgo de lixiviación y pérdida de fertilidad.",
        fill="FBE7EA",
        border=COLORS["red"],
    )
    doc.add_page_break()

    add_title(doc, "6. Reto final: diseña un suelo estable", 1)
    add_body(
        doc,
        "Tu equipo debe encontrar una combinación de variables que mantenga disponibilidad para plantas, fijación biológica y balance del ciclo en niveles altos, con bajo estrés por amoníaco y bajo riesgo de lixiviación.",
    )
    add_colored_table(
        doc,
        ["Criterio", "Meta", "Valor logrado", "Evidencia"],
        [
            ["Disponibilidad para plantas", "Alta", "", ""],
            ["Fijación biológica", "Media-alta", "", ""],
            ["Riesgo de lixiviación", "Bajo o controlado", "", ""],
            ["Estrés por amoníaco", "Bajo", "", ""],
            ["Balance del ciclo", "Alto", "", ""],
        ],
        header_fill=COLORS["navy"],
        band_fill="F2FBFC",
    )
    add_callout(
        doc,
        "Producto final",
        "Escribe una conclusión de 8 a 10 líneas: describe tu combinación, defiende por qué funciona y advierte qué pasaría si una variable se sale de rango.",
        fill="FFF8E8",
        border=COLORS["gold"],
    )
    add_title(doc, "Rúbrica rápida", 2)
    add_colored_table(
        doc,
        ["Criterio", "Excelente", "En proceso"],
        [
            ["Registro de datos", "Completa valores y observaciones claras.", "Faltan datos o son difíciles de interpretar."],
            ["Análisis causal", "Relaciona variables con procesos del ciclo.", "Describe cambios sin explicar la causa."],
            ["Uso de evidencia", "Cita indicadores, barras y consecuencias.", "Usa opiniones sin datos del simulador."],
            ["Conclusión", "Propone un equilibrio y reconoce riesgos.", "No justifica por qué el sistema es estable."],
        ],
        header_fill=COLORS["teal"],
        band_fill="E8F8FA",
    )
    doc.add_paragraph()
    add_body(doc, "Firma del equipo: ____________________________________________", bold=True, color=COLORS["navy"])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
