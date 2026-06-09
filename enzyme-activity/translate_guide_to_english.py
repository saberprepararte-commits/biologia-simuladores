from pathlib import Path
from docx import Document

SRC = Path("Guia_indagacion_Enzymes_Activity_BioLab.docx")
OUT = Path("Inquiry_Guide_Enzymes_Activity_BioLab.docx")

REPLACEMENTS = {
    "Enzymes Activity BioLab | Guía de indagación": "Enzymes Activity BioLab | Inquiry guide",
    "Competencia de indagación: preguntar, experimentar, analizar y argumentar.": "Inquiry competency: asking questions, experimenting, analyzing, and constructing arguments.",
    "Guía de indagación: temperatura y actividad enzimática": "Inquiry guide: temperature and enzymatic activity",
    "Enzymes Activity BioLab | Taller para estudiante": "Enzymes Activity BioLab | Student workshop",
    "Vista de trabajo del simulador: tubos, zona de calor, almidon y panel de resultados.": "Simulator workspace view: tubes, heat zone, starch, and results panel.",
    "1. Pregunta problema": "1. Problem question",
    "Durante la práctica observarás diez tubos con saliva. Algunos se expondrán a calor y otros permanecerán sin calentamiento. Luego todos recibirán almidón y se registrará el cambio observado.": "During the activity, you will observe ten tubes with saliva. Some will be exposed to heat and others will remain unheated. Then all tubes will receive starch, and the observed change will be recorded.",
    "Pregunta investigable que voy a responder:": "Investigable question I will answer:",
    "2. Antes de simular: predicción": "2. Before simulating: prediction",
    "No busques acertar de memoria:": "Do not try to guess from memory:",
    "formula una hipótesis que pueda ponerse a prueba. Lo importante es que conectes variable, resultado esperado y razón biológica.": "write a hypothesis that can be tested. The important thing is to connect the variable, the expected result, and the biological reason.",
    "Hipótesis: Si": "Hypothesis: If",
    "entonces": "then",
    "porque": "because",
    "3. Diseño experimental": "3. Experimental design",
    "Elemento": "Element",
    "Lo que debes identificar en la simulación": "What you should identify in the simulation",
    "Variable independiente": "Independent variable",
    "Variable dependiente": "Dependent variable",
    "Variables controladas": "Controlled variables",
    "Grupo con tratamiento": "Treatment group",
    "Grupo de comparación": "Comparison group",
    "Pista: piensa en qué cambia intencionalmente, qué se mide y qué debe mantenerse igual para comparar de manera justa.": "Hint: think about what is intentionally changed, what is measured, and what must remain the same in order to make a fair comparison.",
    "4. Procedimiento de trabajo": "4. Work procedure",
    "Prepara la práctica en el simulador y observa el estado inicial de los 10 tubos.": "Prepare the activity in the simulator and observe the initial state of the 10 tubes.",
    "Arrastra cinco tubos a la zona de calor. Registra cuáles fueron tratados.": "Drag five tubes to the heat zone. Record which ones were treated.",
    "Agrega almidón a todos los tubos y verifica que los diez queden incluidos.": "Add starch to all tubes and verify that all ten are included.",
    "Inicia el tiempo de reacción y espera hasta obtener los resultados.": "Start the reaction time and wait until the results are obtained.",
    "Registra datos numéricos y observaciones visuales antes de sacar conclusiones.": "Record numerical data and visual observations before drawing conclusions.",
    "5. Registro de datos": "5. Data record",
    "Tubo": "Tube",
    "Grupo / tratamiento": "Group / treatment",
    "Calor": "Heat",
    "Almidón": "Starch",
    "Cambio observado (%)": "Observed change (%)",
    "Observaciones de color o aspecto": "Color or appearance observations",
    "6. Análisis cuantitativo": "6. Quantitative analysis",
    "Cálculo": "Calculation",
    "Operación": "Operation",
    "Resultado": "Result",
    "Promedio del grupo sin calor": "Average for the unheated group",
    "Promedio del grupo con calor": "Average for the heated group",
    "Diferencia entre promedios": "Difference between averages",
    "Relación entre ambos grupos": "Relationship between both groups",
    "Patrón principal que observo en los datos:": "Main pattern I observe in the data:",
    "Dato que más apoya mi explicación:": "Data point that best supports my explanation:",
    "7. Representación gráfica": "7. Graphical representation",
    "Elabora una gráfica de barras con dos categorías: grupo sin calor y grupo con calor. Usa el eje vertical para el cambio promedio observado (%).": "Create a bar graph with two categories: unheated group and heated group. Use the vertical axis for the average observed change (%).",
    "Espacio para gráfica": "Space for graph",
    "8. Interpretación biológica": "8. Biological interpretation",
    "¿Qué explicación biológica es compatible con la diferencia entre grupos?": "What biological explanation is compatible with the difference between groups?",
    "Respuesta:": "Answer:",
    "¿Por qué la amilasa salival y el almidón forman un modelo más realista que saliva y carne para esta pregunta?": "Why do salivary amylase and starch form a more realistic model than saliva and meat for this question?",
    "¿Qué evidencia necesitarías para afirmar que la temperatura cambió la actividad enzimática y no otro factor?": "What evidence would you need to claim that temperature changed enzymatic activity and not another factor?",
    "9. Conclusión con modelo CER": "9. Conclusion using the CER model",
    "Componente": "Component",
    "Escribe tu respuesta": "Write your response",
    "Afirmación": "Claim",
    "¿Qué respondes a la pregunta problema?": "How do you answer the problem question?",
    "Evidencia": "Evidence",
    "¿Qué datos concretos respaldan tu afirmación?": "What specific data support your claim?",
    "Razonamiento": "Reasoning",
    "¿Cómo conectas esos datos con la acción de una enzima?": "How do you connect those data to the action of an enzyme?",
    "10. Extensión investigativa": "10. Inquiry extension",
    "Propón una nueva versión del experimento para obtener datos más complejos.": "Propose a new version of the experiment to obtain more complex data.",
    "Decisión de investigación": "Research decision",
    "Tu propuesta": "Your proposal",
    "Nueva pregunta": "New question",
    "Niveles de temperatura o tiempo que compararías": "Temperature or time levels you would compare",
    "Número de repeticiones por tratamiento": "Number of repetitions per treatment",
    "Dato adicional que medirías": "Additional data you would measure",
    "Cómo esperas que cambie la gráfica": "How you expect the graph to change",
    "Rúbrica rápida": "Quick rubric",
    "Criterio": "Criterion",
    "Logro alto": "High achievement",
    "En proceso": "In progress",
    "Hipótesis": "Hypothesis",
    "Relaciona variable, resultado y razón biológica.": "Connects the variable, result, and biological reason.",
    "Predice sin justificar o sin variable clara.": "Predicts without justification or without a clear variable.",
    "Datos": "Data",
    "Registra los 10 tubos y distingue tratamientos.": "Records the 10 tubes and distinguishes treatments.",
    "Faltan datos o no diferencia grupos.": "Data are missing or groups are not differentiated.",
    "Análisis": "Analysis",
    "Calcula promedios y compara con evidencia.": "Calculates averages and compares using evidence.",
    "Describe sin calcular o sin comparar.": "Describes without calculating or comparing.",
    "Conclusión CER": "CER conclusion",
    "Afirmación, evidencia y razonamiento están conectados.": "Claim, evidence, and reasoning are connected.",
    "La conclusión no usa datos suficientes.": "The conclusion does not use sufficient data.",
    "Indagación": "Inquiry",
    "Propone mejora experimental medible.": "Proposes a measurable experimental improvement.",
    "Propone cambios vagos o no evaluables.": "Proposes vague or non-evaluable changes.",
    "Glosario mínimo": "Minimum glossary",
    "Término": "Term",
    "Definición de trabajo": "Working definition",
    "Enzima": "Enzyme",
    "Proteína que acelera una reacción química sin consumirse en ella.": "Protein that speeds up a chemical reaction without being consumed in it.",
    "Amilasa salival": "Salivary amylase",
    "Enzima presente en la saliva que actúa sobre el almidón.": "Enzyme present in saliva that acts on starch.",
    "Sustrato": "Substrate",
    "Sustancia sobre la que actúa una enzima; en esta práctica, el almidón.": "Substance on which an enzyme acts; in this activity, starch.",
    "Actividad enzimática": "Enzymatic activity",
    "Cambio observable asociado con la acción de la enzima sobre su sustrato.": "Observable change associated with the action of the enzyme on its substrate.",
    "Tratamiento térmico": "Heat treatment",
    "Exposición a calor antes de comparar los resultados con otro grupo.": "Exposure to heat before comparing the results with another group.",
    "Nombre": "Name",
    "Curso / grupo": "Course / group",
    "Fecha": "Date",
    "Integrantes": "Team members",
    "Reto de indagación:": "Inquiry challenge:",
    "usar la simulación para descubrir, con datos, qué ocurre con la actividad de la saliva cuando algunos tubos reciben un tratamiento térmico antes de reaccionar con almidón.": "use the simulation to discover, with data, what happens to saliva activity when some tubes receive heat treatment before reacting with starch.",
}


def translate_run_text(text):
    for source in sorted(REPLACEMENTS, key=len, reverse=True):
        text = text.replace(source, REPLACEMENTS[source])
    return text


def translate_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = translate_run_text(run.text)


def translate_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph(paragraph)
            for nested in cell.tables:
                translate_table(nested)


doc = Document(SRC)

for section in doc.sections:
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            translate_paragraph(paragraph)
        for table in part.tables:
            translate_table(table)

for paragraph in doc.paragraphs:
    translate_paragraph(paragraph)
for table in doc.tables:
    translate_table(table)

doc.save(OUT)
print(OUT.resolve())
