from pathlib import Path
from docx import Document

path = Path("Inquiry_Guide_Enzymes_Activity_BioLab.docx")
doc = Document(path)

target = "Substance on which an enzyme acts; in this activity, starch."

def fix_paragraph(paragraph):
    return

def fix_table(table):
    for row in table.rows:
        if len(row.cells) >= 2 and row.cells[0].text.strip() == "Substrate":
            row.cells[1].text = target
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                fix_paragraph(paragraph)
            for nested in cell.tables:
                fix_table(nested)

for section in doc.sections:
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            fix_paragraph(paragraph)
        for table in part.tables:
            fix_table(table)

for paragraph in doc.paragraphs:
    fix_paragraph(paragraph)
for table in doc.tables:
    fix_table(table)

doc.save(path)
