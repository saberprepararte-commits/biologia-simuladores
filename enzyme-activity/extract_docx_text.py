from pathlib import Path
from docx import Document

doc = Document(Path("Guia_indagacion_Enzymes_Activity_BioLab.docx"))

seen = []

def add_text(text):
    text = text.strip()
    if text and text not in seen:
        seen.append(text)

for section in doc.sections:
    for part in (section.header, section.footer):
        for p in part.paragraphs:
            add_text(p.text)
        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        add_text(p.text)

for p in doc.paragraphs:
    add_text(p.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                add_text(p.text)

for item in seen:
    print(item)
