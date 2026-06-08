from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_student_guide import (
    ASSETS,
    COLORS,
    add_callout,
    add_checkbox_list,
    add_heading,
    add_lines,
    add_table,
    get_font,
    set_cell_text,
    set_styles,
)


BASE = Path(__file__).resolve().parent
OUT = BASE / "Student Inquiry Guide - Cellular Respiration BioLab.docx"
SHOT = BASE / "guide-screenshot-clean-en.png"
MOLECULE_STRIP = BASE / "guide_respiration_molecules_en.png"
INQUIRY_IMG = BASE / "guide_inquiry_cycle_en.png"


def create_molecule_strip_en():
    items = [
        ("molecule-glucose-ai.png", "Glucose", "C6H12O6"),
        ("molecule-o2-ai.png", "Oxygen", "O2"),
        ("molecule-pyruvate-ai.png", "Pyruvate", "C3H3O3-"),
        ("molecule-atp-ai.png", "ATP", "energy"),
        ("molecule-co2-ai.png", "Carbon dioxide", "CO2"),
        ("molecule-h2o-ai.png", "Water", "H2O"),
        ("molecule-lactate-ai.png", "Lactate", "C3H5O3-"),
    ]
    w, h = 1500, 310
    img = Image.new("RGBA", (w, h), "#EAF8FB")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, True)
    label_font = get_font(23, True)
    small_font = get_font(20, False)
    draw.rounded_rectangle((14, 14, w - 14, h - 14), radius=34, fill="#EAF8FB", outline="#74C8D8", width=4)
    draw.text((42, 28), "Molecular nomenclature for interpreting the model", fill="#0B3148", font=title_font)
    cell_w = (w - 84) // len(items)
    for i, (file_name, label, formula) in enumerate(items):
        x = 42 + i * cell_w
        cx = x + cell_w // 2
        if i:
            draw.line((x - 12, 96, x - 12, h - 32), fill="#C7E5EC", width=2)
        mol = Image.open(ASSETS / file_name).convert("RGBA")
        mol.thumbnail((128, 112), Image.LANCZOS)
        img.alpha_composite(mol, (cx - mol.width // 2, 100))
        tw = draw.textbbox((0, 0), label, font=label_font)
        draw.text((cx - (tw[2] - tw[0]) / 2, 220), label, fill="#0B3148", font=label_font)
        fw = draw.textbbox((0, 0), formula, font=small_font)
        draw.text((cx - (fw[2] - fw[0]) / 2, 252), formula, fill="#4B6172", font=small_font)
    img.convert("RGB").save(MOLECULE_STRIP, quality=95)


def create_inquiry_image_en():
    w, h = 1400, 430
    img = Image.new("RGBA", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = get_font(42, True)
    label_font = get_font(26, True)
    small_font = get_font(22, False)
    draw.rounded_rectangle((18, 18, w - 18, h - 18), radius=34, fill="#F6FBFD", outline="#91D7E5", width=4)
    draw.text((54, 36), "Inquiry pathway with the simulator", fill="#0B3148", font=title_font)
    nodes = [
        ("Question", "Which variable changes ATP the most?", "#BDEFF7"),
        ("Hypothesis", "Explained prediction", "#FFF1BE"),
        ("Experiment", "Change one variable", "#DFF7E8"),
        ("Record", "Data and observations", "#ECE0FF"),
        ("Argue", "Evidence-based conclusion", "#F7D8DF"),
    ]
    x0, y0, gap = 64, 165, 260
    for i, (title, text, fill) in enumerate(nodes):
        x = x0 + i * gap
        draw.rounded_rectangle((x, y0, x + 210, y0 + 150), radius=26, fill=fill, outline="#7EAEC0", width=3)
        draw.text((x + 22, y0 + 24), title, fill="#0B3148", font=label_font)
        draw.multiline_text((x + 22, y0 + 72), text, fill="#36576A", font=small_font, spacing=4)
        if i < len(nodes) - 1:
            draw.line((x + 214, y0 + 75, x + gap - 10, y0 + 75), fill="#2E9BC0", width=8)
            draw.polygon([(x + gap - 10, y0 + 75), (x + gap - 35, y0 + 60), (x + gap - 35, y0 + 90)], fill="#2E9BC0")
    img.convert("RGB").save(INQUIRY_IMG, quality=95)


def add_cover_en(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Inquiry Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("Cellular Respiration BioLab")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Virtual laboratory for developing skills in questioning, hypothesis building, variable control, data analysis, and scientific argumentation."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.add_picture(str(SHOT), width=Inches(6.45))
    add_callout(
        doc,
        "Central challenge",
        "Use simulator data to explain how glucose availability, oxygen, mitochondrial activity, energy demand, and temperature modify the production of ATP, CO2, water, heat, and lactate.",
        COLORS["gold"],
    )
    table = add_table(
        doc,
        ["Student", "Class", "Date", "Group"],
        [["", "", "", ""]],
        [1.9, 1.4, 1.4, 1.5],
        COLORS["cyan"],
        9,
    )
    for cell in table.rows[1].cells:
        set_cell_text(cell, "\n", size=9)


def build_doc_en():
    create_molecule_strip_en()
    create_inquiry_image_en()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Cellular Respiration BioLab - Student Inquiry Guide")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_cover_en(doc)
    doc.add_page_break()

    add_heading(doc, "1. Before experimenting: scientific orientation", 1)
    doc.add_picture(str(INQUIRY_IMG), width=Inches(6.45))
    add_callout(
        doc,
        "Key idea",
        "A simulator is not used to guess answers: it is used to test relationships. Change one variable at a time, keep a control group, and support your conclusions with numerical and observational evidence.",
        COLORS["cyan2"],
    )
    add_heading(doc, "Learning objectives", 2)
    add_checkbox_list(
        doc,
        [
            "Formulate an investigable question about cellular respiration.",
            "Design a fair test by changing one independent variable and keeping the others constant.",
            "Record ATP/glucose, aerobic efficiency, oxygen debt, and lactate data.",
            "Explain why the cell shifts between aerobic respiration and lactic fermentation.",
            "Build a conclusion that connects evidence, biological reasoning, and model limitations.",
        ],
    )
    add_heading(doc, "Simulator variables", 2)
    add_table(
        doc,
        ["Manipulated variable", "Simulator unit", "What it represents biologically", "Initial prediction"],
        [
            ["Available glucose", "mg/dL", "Fuel that enters glycolysis and allows pyruvate formation.", ""],
            ["Available oxygen", "% O2", "Final electron acceptor in the electron transport chain.", ""],
            ["Mitochondrial activity", "%", "Capacity of inner membranes, enzymes, and organelles to produce ATP.", ""],
            ["Energy demand", "%", "The cell's ATP requirement: rest, exercise, or high activity.", ""],
            ["Cell temperature", "degrees C", "Enzyme speed and protein/membrane stability.", ""],
        ],
        [1.55, 1.05, 2.85, 1.05],
        COLORS["gray"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "2. Nomenclature and model reading", 1)
    doc.add_picture(str(MOLECULE_STRIP), width=Inches(6.45))
    add_heading(doc, "Simulator lab readings", 2)
    add_table(
        doc,
        ["Indicator", "Unit", "Reference value for interpretation", "How to use it in your analysis"],
        [
            ["ATP production", "ATP/glucose", "Aerobic respiration is commonly approximated at 30-32 ATP per glucose; the simulator scales up to 32.", "Compare whether your changes move the cell toward or away from efficient production."],
            ["Aerobic efficiency", "%", "High when glucose, O2, and functioning mitochondria are available.", "Explain whether the cell is using its substrates efficiently or whether a limitation appears."],
            ["Oxygen debt", "% O2", "Increases when demand exceeds oxygen availability.", "Relate high oxygen debt to lower aerobic ATP and higher lactate."],
            ["Lactic fermentation", "mmol/L", "Approximate normal blood lactate: 0.5-2.2 mmol/L; higher values suggest accumulation.", "Interpret whether the scenario simulates metabolic stress or low O2 availability."],
        ],
        [1.55, 1.0, 2.25, 1.7],
        COLORS["mint"],
        8,
    )
    add_heading(doc, "Your investigable question", 2)
    doc.add_paragraph("Write a question that can be answered by modifying only one simulator variable.")
    add_lines(doc, 3)
    add_heading(doc, "Hypothesis", 2)
    doc.add_paragraph("Write a causal prediction: if I modify ___, then ___, because ___.")
    add_lines(doc, 4)

    doc.add_page_break()
    add_heading(doc, "3. Experiment A: control group and glucose change", 1)
    add_callout(
        doc,
        "Experimental rule",
        "First press Control group. Then modify only glucose. Do not change oxygen, mitochondrial activity, demand, or temperature during this experiment.",
        COLORS["mint"],
    )
    add_table(
        doc,
        ["Trial", "Glucose (mg/dL)", "ATP/glucose", "Efficiency (%)", "CO2", "Lactate (mmol/L)", "Observation"],
        [
            ["Control", "90", "", "", "", "", ""],
            ["Low glucose", "20-50", "", "", "", "", ""],
            ["Medium glucose", "90", "", "", "", "", ""],
            ["High glucose", "150-180", "", "", "", "", ""],
        ],
        [0.72, 1.05, 0.88, 0.82, 0.62, 1.0, 1.45],
        COLORS["cyan"],
        7,
    )
    add_heading(doc, "Experiment A analysis", 2)
    add_table(
        doc,
        ["Analysis question", "Evidence-based answer"],
        [
            ["What happened to ATP when glucose was very low?", ""],
            ["Did high glucose always increase efficiency? Explain.", ""],
            ["Which variable did you keep constant to make the comparison fair?", ""],
        ],
        [2.15, 4.15],
        COLORS["gray"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "4. Experiment B: oxygen and lactic fermentation", 1)
    add_callout(
        doc,
        "Biological focus",
        "The electron transport chain needs oxygen as its final acceptor. When oxygen decreases, pyruvate may be redirected toward lactate and ATP production falls.",
        COLORS["rose"],
    )
    add_table(
        doc,
        ["Trial", "Oxygen (%)", "ATP/glucose", "O2 debt (%)", "Lactate (mmol/L)", "Oxygen state", "Interpretation"],
        [
            ["Control", "96", "", "", "", "", ""],
            ["Limited", "70", "", "", "", "", ""],
            ["Low", "55", "", "", "", "", ""],
            ["Critical", "30", "", "", "", "", ""],
        ],
        [0.72, 0.85, 0.9, 0.9, 1.0, 1.05, 1.05],
        COLORS["rose"],
        7,
    )
    add_heading(doc, "Argue", 2)
    doc.add_paragraph("Use your data to explain why lactate can increase even when glucose is still available.")
    add_lines(doc, 6)

    doc.add_page_break()
    add_heading(doc, "5. Experiment C: mitochondrion, demand, and temperature", 1)
    add_table(
        doc,
        ["Scenario", "Mitochondrial activity (%)", "Demand (%)", "Temperature (C)", "ATP/glucose", "Cell balance", "Biological explanation"],
        [
            ["Efficient mitochondrion", "80-100", "55", "37", "", "", ""],
            ["Low mitochondrial activity", "10-40", "55", "37", "", "", ""],
            ["High demand", "80", "120-140", "37", "", "", ""],
            ["Heat stress", "80", "55", "40-44", "", "", ""],
        ],
        [1.0, 1.1, 0.78, 0.78, 0.78, 0.85, 1.35],
        COLORS["lavender"],
        7,
    )
    add_heading(doc, "Comparative mini-conclusion", 2)
    doc.add_paragraph("Which factor reduced cell balance the most: low mitochondrial activity, high demand, or extreme temperature? Justify with data.")
    add_lines(doc, 7)

    doc.add_page_break()
    add_heading(doc, "6. Free inquiry design", 1)
    add_callout(
        doc,
        "Your challenge",
        "Design your own experiment with at least three trials. You must control variables, record data, and defend a conclusion.",
        COLORS["gold"],
    )
    add_table(
        doc,
        ["Design element", "Your proposal"],
        [
            ["Investigable question", ""],
            ["Hypothesis", ""],
            ["Independent variable", ""],
            ["Controlled variables", ""],
            ["Main dependent variable", ""],
            ["Criterion for accepting or rejecting the hypothesis", ""],
        ],
        [2.05, 4.25],
        COLORS["gray"],
        8,
    )
    add_table(
        doc,
        ["Trial", "Changed variable", "ATP/glucose", "Efficiency", "O2 debt", "Lactate", "Visual evidence observed"],
        [["1", "", "", "", "", "", ""], ["2", "", "", "", "", "", ""], ["3", "", "", "", "", "", ""], ["4 optional", "", "", "", "", "", ""]],
        [0.62, 1.25, 0.78, 0.78, 0.78, 0.78, 1.35],
        COLORS["cyan"],
        7,
    )
    add_heading(doc, "CER conclusion", 2)
    add_table(
        doc,
        ["C", "Claim", ""],
        [["E", "Evidence", ""], ["R", "Biological reasoning", ""]],
        [0.45, 1.75, 4.1],
        COLORS["mint"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "7. Discussion, transfer, and metacognition", 1)
    add_heading(doc, "Closing questions", 2)
    closing = [
        "Why does the cell not always produce the same amount of ATP even when it has glucose?",
        "What is the difference between having a lot of fuel and having the real capacity to transform it into ATP?",
        "What simulator evidence indicates that the cell is relying more on lactic fermentation?",
        "What limitations does this model have compared with a real cell?",
    ]
    for question in closing:
        doc.add_paragraph(question, style="List Number")
        add_lines(doc, 2)
    add_heading(doc, "Quick rubric", 2)
    add_table(
        doc,
        ["Criterion", "4 - Advanced", "3 - Achieved", "2 - Developing", "1 - Beginning"],
        [
            ["Question and hypothesis", "Measurable question and causal hypothesis.", "Clear question and appropriate hypothesis.", "Broad question or incomplete hypothesis.", "No causal relationship is identified."],
            ["Variable control", "Changes one variable and justifies controls.", "Keeps the main controls constant.", "Changes several variables without explaining.", "No experimental control."],
            ["Data analysis", "Uses numbers, trends, and comparisons.", "Uses sufficient data.", "Mentions data without interpreting them.", "Answers without evidence."],
            ["Biological reasoning", "Connects glycolysis, mitochondrion, O2, ATP, and lactate.", "Explains the main relationship.", "Partial explanation.", "Confuses key processes."],
        ],
        [1.25, 1.3, 1.3, 1.25, 1.2],
        COLORS["orange"],
        7,
    )
    add_heading(doc, "Reference sources", 2)
    refs = [
        "OpenStax Biology 2e, Chapter 7: glycolysis, cellular respiration, and electron transport chain. https://openstax.org/books/biology-2e/pages/7-introduction",
        "MedlinePlus, Lactic Acid Test: approximate normal range 0.5-2.2 mmol/L. https://medlineplus.gov/ency/article/003507.htm",
        "Cleveland Clinic, Blood Oxygen Level: usual normal saturation 95-100%. https://my.clevelandclinic.org/health/diagnostics/22447-blood-oxygen-level",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(ref)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc_en())
