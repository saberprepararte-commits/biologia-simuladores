from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
ASSETS = BASE / "assets"
OUT = BASE / "Guia de indagacion - Respiracion celular BioLab.docx"
SHOT = BASE / "guide-screenshot-clean.png"
MOLECULE_STRIP = BASE / "guia_moleculas_respiracion.png"
INQUIRY_IMG = BASE / "guia_ciclo_indagacion.png"


COLORS = {
    "navy": "0B3148",
    "blue": "2E74B5",
    "deep_blue": "1F4D78",
    "cyan": "BDEFF7",
    "cyan2": "DDF8FC",
    "mint": "DFF7E8",
    "gold": "FFF1BE",
    "orange": "F9D7AE",
    "rose": "F7D8DF",
    "lavender": "ECE0FF",
    "gray": "E8EEF5",
    "light": "F6FBFD",
    "white": "FFFFFF",
    "ink": "12324A",
    "muted": "4B6172",
}


def get_font(size=36, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def create_molecule_strip():
    items = [
        ("molecule-glucose-ai.png", "Glucosa", "C6H12O6"),
        ("molecule-o2-ai.png", "Oxigeno", "O2"),
        ("molecule-pyruvate-ai.png", "Piruvato", "C3H3O3-"),
        ("molecule-atp-ai.png", "ATP", "energia"),
        ("molecule-co2-ai.png", "Dioxido de carbono", "CO2"),
        ("molecule-h2o-ai.png", "Agua", "H2O"),
        ("molecule-lactate-ai.png", "Lactato", "C3H5O3-"),
    ]
    w, h = 1500, 310
    img = Image.new("RGBA", (w, h), "#EAF8FB")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, True)
    label_font = get_font(23, True)
    small_font = get_font(20, False)
    draw.rounded_rectangle((14, 14, w - 14, h - 14), radius=34, fill="#EAF8FB", outline="#74C8D8", width=4)
    draw.text((42, 28), "Nomenclatura molecular para interpretar el modelo", fill="#0B3148", font=title_font)
    cell_w = (w - 84) // len(items)
    for i, (file_name, label, formula) in enumerate(items):
        x = 42 + i * cell_w
        cx = x + cell_w // 2
        if i:
            draw.line((x - 12, 96, x - 12, h - 32), fill="#C7E5EC", width=2)
        mol = Image.open(ASSETS / file_name).convert("RGBA")
        mol.thumbnail((128, 112), Image.LANCZOS)
        mx = cx - mol.width // 2
        img.alpha_composite(mol, (mx, 100))
        tw = draw.textbbox((0, 0), label, font=label_font)
        draw.text((cx - (tw[2] - tw[0]) / 2, 220), label, fill="#0B3148", font=label_font)
        fw = draw.textbbox((0, 0), formula, font=small_font)
        draw.text((cx - (fw[2] - fw[0]) / 2, 252), formula, fill="#4B6172", font=small_font)
    img.convert("RGB").save(MOLECULE_STRIP, quality=95)


def create_inquiry_image():
    w, h = 1400, 430
    img = Image.new("RGBA", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = get_font(42, True)
    label_font = get_font(26, True)
    small_font = get_font(22, False)
    draw.rounded_rectangle((18, 18, w - 18, h - 18), radius=34, fill="#F6FBFD", outline="#91D7E5", width=4)
    draw.text((54, 36), "Ruta de indagacion con el simulador", fill="#0B3148", font=title_font)
    nodes = [
        ("Pregunta", "Que variable cambia mas el ATP?", "#BDEFF7"),
        ("Hipotesis", "Prediccion explicada", "#FFF1BE"),
        ("Experimenta", "Modifica una variable", "#DFF7E8"),
        ("Registra", "Datos y observaciones", "#ECE0FF"),
        ("Argumenta", "Conclusion con evidencia", "#F7D8DF"),
    ]
    x0, y0, gap = 64, 165, 260
    for i, (title, text, fill) in enumerate(nodes):
        x = x0 + i * gap
        draw.rounded_rectangle((x, y0, x + 210, y0 + 150), radius=26, fill=fill, outline="#7EAEC0", width=3)
        draw.text((x + 22, y0 + 24), title, fill="#0B3148", font=label_font)
        draw.multiline_text((x + 22, y0 + 72), text, fill="#36576A", font=small_font, spacing=4)
        if i < len(nodes) - 1:
            draw.line((x + 214, y0 + 75, x + gap - 10, y0 + 75), fill="#2E9BC0", width=8)
            draw.polygon([(x + gap - 10, y0 + 75), (x + gap - 35, y0 + 60), (x + gap - 35, y0 + 90)], fill="#2E9BC0")
    img.convert("RGB").save(INQUIRY_IMG, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="12324A", size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def shade_header(row, fill=COLORS["gray"]):
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(COLORS["navy"])


def add_table(doc, headers, rows, widths, header_fill=COLORS["gray"], font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            set_cell_text(row.cells[i], text, size=font_size)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_callout(doc, title, body, fill=COLORS["cyan2"]):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def add_lines(doc, count=3):
    for _ in range(count):
        p = doc.add_paragraph("_" * 92)
        p.paragraph_format.space_after = Pt(2)


def add_checkbox_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(10)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["deep_blue"], 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def apply_spanish_orthography(doc):
    replacements = {
        "Guia": "Guía",
        "indagacion": "indagación",
        "orientacion": "orientación",
        "Respiracion": "Respiración",
        "respiracion": "respiración",
        "celular": "celular",
        "hipotesis": "hipótesis",
        "Hipotesis": "Hipótesis",
        "analisis": "análisis",
        "Analisis": "Análisis",
        "argumentacion": "argumentación",
        "cientifica": "científica",
        "cientifico": "científico",
        "cientificas": "científicas",
        "proposito": "propósito",
        "energetica": "energética",
        "Energetica": "Energética",
        "Oxigeno": "Oxígeno",
        "oxigeno": "oxígeno",
        "Dioxido": "Dióxido",
        "Carbono": "carbono",
        "Agua": "Agua",
        "Glucolisis": "Glucólisis",
        "glucolisis": "glucólisis",
        "biologico": "biológico",
        "biologicamente": "biológicamente",
        "aerobica": "aeróbica",
        "aerobico": "aeróbico",
        "aerobica": "aeróbica",
        "enzimatica": "enzimática",
        "enzimatico": "enzimático",
        "enzimas": "enzimas",
        "mitocondrial": "mitocondrial",
        "celula": "célula",
        "Celula": "Célula",
        "molecula": "molécula",
        "moleculas": "moléculas",
        "quimica": "química",
        "quimico": "químico",
        "aproximo": "aproximó",
        "aproximado": "aproximado",
        "produccion": "producción",
        "Produccion": "Producción",
        "prediccion": "predicción",
        "Prediccion": "Predicción",
        "oxidacion": "oxidación",
        "relacion": "relación",
        "sintesis": "síntesis",
        "fermentacion": "fermentación",
        "lactica": "láctica",
        "energetico": "energético",
        "energetica": "energética",
        "comparacion": "comparación",
        "conclusion": "conclusión",
        "Conclusion": "Conclusión",
        "Discusion": "Discusión",
        "metacognicion": "metacognición",
        "demas": "demás",
        "capitulo": "capítulo",
        "rubrica": "rúbrica",
        "Rapida": "Rápida",
        "rapida": "rápida",
        "diseno": "diseño",
        "Diseno": "Diseño",
        "que ": "qué ",
        "Que ": "Qué ",
        "Como ": "Cómo ",
        "como ": "cómo ",
        "Cual ": "Cuál ",
        "cual ": "cuál ",
        "por que": "por qué",
        "Por que": "Por qué",
        "alli": "allí",
        "ahi": "ahí",
        "mas ": "más ",
        "limites": "límites",
        "numero": "número",
        "minimo": "mínimo",
        "detras": "detrás",
        "todavia": "todavía",
        "esta ": "está ",
        "Esta ": "Está ",
        "si ": "si ",
        "C ": "°C ",
        "grados °C": "°C",
        "célular": "celular",
        "Célular": "Celular",
        "conclusión qué conecte": "conclusión que conecte",
        "pregunta qué pueda": "pregunta que pueda",
        "con datos del simulador, cómo": "con datos del simulador, cómo",
        "porqué ___": "porque ___",
        "aunqué": "aunque",
        "Que variable": "Qué variable",
        "Que representa": "Qué representa",
        "Que ocurrió": "Qué ocurrió",
        "Que ocurrio": "Qué ocurrió",
        "Que diferencia": "Qué diferencia",
        "Que evidencia": "Qué evidencia",
        "Que limitaciones": "Qué limitaciones",
        "Qué puede": "Que puede",
        "que pueda": "que pueda",
        "que conecte": "que conecte",
        "relaciónes": "relaciones",
        "conclusiónes": "conclusiones",
        "Combustible qué ingresa": "Combustible que ingresa",
        "enzimás": "enzimas",
        "proteinas": "proteínas",
        "investigacion": "investigación",
        "Rubrica": "Rúbrica",
        "saturacion": "saturación",
        "limitacion": "limitación",
        "indica qué la célula": "indica que la célula",
        "grados C": "°C",
        "Cuál factor redujo": "¿Cuál factor redujo",
        "Por qué la célula": "¿Por qué la célula",
        "Qué diferencia hay": "¿Qué diferencia hay",
        "Qué evidencia del simulador": "¿Qué evidencia del simulador",
        "Qué limitaciones tiene": "¿Qué limitaciones tiene",
        "glucosa?": "glucosa?",
        "ATP?": "ATP?",
        "láctica?": "láctica?",
        "real?": "real?",
        "limitaciónes": "limitaciones",
        "Qué limitaciones tiene": "¿Qué limitaciones tiene",
        "Fermentacion": "Fermentación",
        "sanguineo": "sanguíneo",
        "acumulacion": "acumulación",
        "estres": "estrés",
        "metabolico": "metabólico",
        "Despues": "Después",
        "despues": "después",
        "Observacion": "Observación",
        "Respuestá": "Respuesta",
        "Qué ocurrió con": "¿Qué ocurrió con",
        "Qué ocurrio con": "¿Qué ocurrió con",
        "siempre aumento": "siempre aumentó",
        "para qué la comparación": "para que la comparación",
        "cómo aceptor": "como aceptor",
        "Interpretacion": "Interpretación",
        "Critico": "Crítico",
        "Explicacion biologica": "Explicación biológica",
        "Temperatura (C)": "Temperatura (°C)",
        "Qué variable mantuviste": "¿Qué variable mantuviste",
        "La glucosa alta siempre aumentó": "¿La glucosa alta siempre aumentó",
        "Estres termico": "Estrés térmico",
        "afirmacion": "afirmación",
        "comparaciónes": "comparaciones",
        "Explicacion parcial": "Explicación parcial",
    }

    def fix_text(text):
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    def fix_paragraph(paragraph):
        for run in paragraph.runs:
            run.text = fix_text(run.text)

    for paragraph in doc.paragraphs:
        fix_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            fix_paragraph(paragraph)
        for paragraph in section.header.paragraphs:
            fix_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_paragraph(paragraph)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guia de indagacion")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("Respiracion celular BioLab")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Laboratorio virtual para desarrollar habilidades de pregunta, hipotesis, control de variables, analisis de datos y argumentacion cientifica.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    doc.add_picture(str(SHOT), width=Inches(6.45))
    add_callout(
        doc,
        "Reto central",
        "Explica, con datos del simulador, como la disponibilidad de glucosa, oxigeno, actividad mitocondrial, demanda energetica y temperatura modifican la produccion de ATP, CO2, agua, calor y lactato.",
        COLORS["gold"],
    )
    table = add_table(
        doc,
        ["Estudiante", "Curso", "Fecha", "Grupo"],
        [["", "", "", ""]],
        [1.9, 1.4, 1.4, 1.5],
        COLORS["cyan"],
        9,
    )
    for c in table.rows[1].cells:
        set_cell_text(c, "\n", size=9)


def build_doc():
    create_molecule_strip()
    create_inquiry_image()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Respiracion celular BioLab - Guia de indagacion estudiantil")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_cover(doc)
    doc.add_page_break()

    add_heading(doc, "1. Antes de experimentar: orientacion cientifica", 1)
    doc.add_picture(str(INQUIRY_IMG), width=Inches(6.45))
    add_callout(
        doc,
        "Idea clave",
        "Un simulador no se usa para adivinar respuestas: se usa para probar relaciones. Cambia una variable a la vez, conserva un grupo control y argumenta tus conclusiones con evidencia numerica y observacional.",
        COLORS["cyan2"],
    )
    add_heading(doc, "Objetivos de aprendizaje", 2)
    add_checkbox_list(
        doc,
        [
            "Formular una pregunta investigable sobre respiracion celular.",
            "Diseñar una prueba justa modificando una variable independiente y manteniendo constantes las demas.",
            "Registrar datos de ATP/glucosa, eficiencia aerobica, deuda de oxigeno y lactato.",
            "Explicar por que la celula cambia entre respiracion aerobica y fermentacion lactica.",
            "Construir una conclusion que conecte evidencia, razonamiento biologico y limites del modelo.",
        ],
    )
    add_heading(doc, "Variables del simulador", 2)
    add_table(
        doc,
        ["Variable manipulable", "Unidad en el simulador", "Que representa biologicamente", "Prediccion inicial"],
        [
            ["Glucosa disponible", "mg/dL", "Combustible que ingresa a glucolisis y permite formar piruvato.", ""],
            ["Oxigeno disponible", "% O2", "Aceptor final de electrones en la cadena de transporte.", ""],
            ["Actividad mitocondrial", "%", "Capacidad de membranas internas, enzimas y organelos para producir ATP.", ""],
            ["Demanda energetica", "%", "Necesidad de ATP de la celula: reposo, ejercicio o alta actividad.", ""],
            ["Temperatura celular", "grados C", "Velocidad enzimatica y estabilidad de proteinas/membranas.", ""],
        ],
        [1.55, 1.05, 2.85, 1.05],
        COLORS["gray"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "2. Nomenclatura y lectura del modelo", 1)
    doc.add_picture(str(MOLECULE_STRIP), width=Inches(6.45))
    add_heading(doc, "Lecturas de laboratorio del simulador", 2)
    add_table(
        doc,
        ["Indicador", "Unidad", "Valor de referencia para interpretar", "Como usarlo en tu analisis"],
        [
            ["Produccion de ATP", "ATP/glucosa", "Respiracion aerobica suele aproximarse a 30-32 ATP por glucosa; el simulador escala hasta 32.", "Compara si tus cambios acercan o alejan la celula de una produccion eficiente."],
            ["Eficiencia aerobica", "%", "Alta cuando hay glucosa, O2 y mitocondrias funcionales.", "Explica si la celula usa bien sus sustratos o si hay una limitacion."],
            ["Deuda de oxigeno", "% O2", "Aumenta cuando la demanda supera la disponibilidad de oxigeno.", "Relaciona deuda alta con menor ATP aerobico y mayor lactato."],
            ["Fermentacion lactica", "mmol/L", "Lactato sanguineo normal aproximado: 0.5-2.2 mmol/L; valores mayores sugieren acumulacion.", "Interpreta si el escenario simula estres metabolico o baja disponibilidad de O2."],
        ],
        [1.55, 1.0, 2.25, 1.7],
        COLORS["mint"],
        8,
    )
    add_heading(doc, "Pregunta investigable propia", 2)
    doc.add_paragraph("Escribe una pregunta que pueda responderse modificando una sola variable del simulador.")
    add_lines(doc, 3)
    add_heading(doc, "Hipotesis", 2)
    doc.add_paragraph("Redacta una prediccion causal: si modifico ___, entonces ___, porque ___.")
    add_lines(doc, 4)

    doc.add_page_break()
    add_heading(doc, "3. Experimento A: grupo control y cambio de glucosa", 1)
    add_callout(
        doc,
        "Regla experimental",
        "Primero presiona Grupo control. Despues modifica solo la glucosa. No cambies oxigeno, actividad mitocondrial, demanda ni temperatura durante este experimento.",
        COLORS["mint"],
    )
    add_table(
        doc,
        ["Ensayo", "Glucosa (mg/dL)", "ATP/glucosa", "Eficiencia (%)", "CO2", "Lactato (mmol/L)", "Observacion"],
        [
            ["Control", "90", "", "", "", "", ""],
            ["Glucosa baja", "20-50", "", "", "", "", ""],
            ["Glucosa media", "90", "", "", "", "", ""],
            ["Glucosa alta", "150-180", "", "", "", "", ""],
        ],
        [0.72, 1.05, 0.88, 0.82, 0.62, 1.0, 1.45],
        COLORS["cyan"],
        7,
    )
    add_heading(doc, "Analisis del experimento A", 2)
    add_table(
        doc,
        ["Pregunta de analisis", "Respuesta con evidencia"],
        [
            ["Que ocurrio con el ATP cuando la glucosa fue muy baja?", ""],
            ["La glucosa alta siempre aumento la eficiencia? Explica.", ""],
            ["Que variable mantuviste constante para que la comparacion fuera justa?", ""],
        ],
        [2.15, 4.15],
        COLORS["gray"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "4. Experimento B: oxigeno y fermentacion lactica", 1)
    add_callout(
        doc,
        "Foco biologico",
        "La cadena de transporte de electrones necesita oxigeno como aceptor final. Cuando el oxigeno baja, el piruvato puede desviarse hacia lactato y la produccion de ATP cae.",
        COLORS["rose"],
    )
    add_table(
        doc,
        ["Ensayo", "Oxigeno (%)", "ATP/glucosa", "Deuda O2 (%)", "Lactato (mmol/L)", "Estado del oxigeno", "Interpretacion"],
        [
            ["Control", "96", "", "", "", "", ""],
            ["Limitado", "70", "", "", "", "", ""],
            ["Bajo", "55", "", "", "", "", ""],
            ["Critico", "30", "", "", "", "", ""],
        ],
        [0.72, 0.85, 0.9, 0.9, 1.0, 1.05, 1.05],
        COLORS["rose"],
        7,
    )
    add_heading(doc, "Argumenta", 2)
    doc.add_paragraph("Usa tus datos para explicar por que puede aumentar el lactato aunque todavia haya glucosa disponible.")
    add_lines(doc, 6)

    doc.add_page_break()
    add_heading(doc, "5. Experimento C: mitocondria, demanda y temperatura", 1)
    add_table(
        doc,
        ["Escenario", "Actividad mitocondrial (%)", "Demanda (%)", "Temperatura (C)", "ATP/glucosa", "Balance celular", "Explicacion biologica"],
        [
            ["Mitocondria eficiente", "80-100", "55", "37", "", "", ""],
            ["Baja actividad mitocondrial", "10-40", "55", "37", "", "", ""],
            ["Alta demanda", "80", "120-140", "37", "", "", ""],
            ["Estres termico", "80", "55", "40-44", "", "", ""],
        ],
        [1.0, 1.1, 0.78, 0.78, 0.78, 0.85, 1.35],
        COLORS["lavender"],
        7,
    )
    add_heading(doc, "Mini-conclusion comparativa", 2)
    doc.add_paragraph("Cual factor redujo mas el balance celular: baja actividad mitocondrial, alta demanda o temperatura extrema? Justifica con datos.")
    add_lines(doc, 7)

    doc.add_page_break()
    add_heading(doc, "6. Diseno de investigacion libre", 1)
    add_callout(
        doc,
        "Tu reto",
        "Diseña un experimento propio con minimo tres ensayos. Debes controlar variables, registrar datos y defender una conclusion.",
        COLORS["gold"],
    )
    add_table(
        doc,
        ["Elemento del diseno", "Tu propuesta"],
        [
            ["Pregunta investigable", ""],
            ["Hipotesis", ""],
            ["Variable independiente", ""],
            ["Variables controladas", ""],
            ["Variable dependiente principal", ""],
            ["Criterio para aceptar o rechazar la hipotesis", ""],
        ],
        [2.05, 4.25],
        COLORS["gray"],
        8,
    )
    add_table(
        doc,
        ["Ensayo", "Variable cambiada", "ATP/glucosa", "Eficiencia", "Deuda O2", "Lactato", "Evidencia visual observada"],
        [["1", "", "", "", "", "", ""], ["2", "", "", "", "", "", ""], ["3", "", "", "", "", "", ""], ["4 opcional", "", "", "", "", "", ""]],
        [0.62, 1.25, 0.78, 0.78, 0.78, 0.78, 1.35],
        COLORS["cyan"],
        7,
    )
    add_heading(doc, "Conclusion CER", 2)
    add_table(
        doc,
        ["C", "Claim / afirmacion", ""],
        [["E", "Evidence / evidencia", ""], ["R", "Reasoning / razonamiento biologico", ""]],
        [0.45, 1.75, 4.1],
        COLORS["mint"],
        8,
    )

    doc.add_page_break()
    add_heading(doc, "7. Discusion, transferencia y metacognicion", 1)
    add_heading(doc, "Preguntas de cierre", 2)
    closing = [
        "Por que la celula no produce siempre la misma cantidad de ATP aunque tenga glucosa?",
        "Que diferencia hay entre tener mucho combustible y tener capacidad real de transformarlo en ATP?",
        "Que evidencia del simulador indica que la celula esta recurriendo mas a fermentacion lactica?",
        "Que limitaciones tiene este modelo frente a una celula real?",
    ]
    for q in closing:
        doc.add_paragraph(q, style="List Number")
        add_lines(doc, 2)
    add_heading(doc, "Rubrica rapida", 2)
    add_table(
        doc,
        ["Criterio", "4 - Avanzado", "3 - Logrado", "2 - En proceso", "1 - Inicial"],
        [
            ["Pregunta e hipotesis", "Pregunta medible e hipotesis causal.", "Pregunta clara e hipotesis adecuada.", "Pregunta amplia o hipotesis incompleta.", "No se identifica relacion causal."],
            ["Control de variables", "Cambia una variable y justifica controles.", "Mantiene controles principales.", "Cambia varias variables sin explicar.", "No hay control experimental."],
            ["Analisis de datos", "Usa numeros, tendencias y comparaciones.", "Usa datos suficientes.", "Menciona datos sin interpretarlos.", "Responde sin evidencia."],
            ["Razonamiento biologico", "Conecta glucolisis, mitocondria, O2, ATP y lactato.", "Explica la relacion principal.", "Explicacion parcial.", "Confunde procesos clave."],
        ],
        [1.25, 1.3, 1.3, 1.25, 1.2],
        COLORS["orange"],
        7,
    )
    add_heading(doc, "Fuentes de referencia", 2)
    refs = [
        "OpenStax Biology 2e, capitulo 7: glucolisis, respiracion celular y cadena de transporte de electrones. https://openstax.org/books/biology-2e/pages/7-introduction",
        "MedlinePlus, Lactic Acid Test: rango normal aproximado 0.5-2.2 mmol/L. https://medlineplus.gov/ency/article/003507.htm",
        "Cleveland Clinic, Blood Oxygen Level: saturacion normal habitual 95-100%. https://my.clevelandclinic.org/health/diagnostics/22447-blood-oxygen-level",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(ref)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    apply_spanish_orthography(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
